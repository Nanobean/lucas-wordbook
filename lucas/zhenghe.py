# -*- coding: utf-8 -*-
"""
Created on Mon Dec 10 13:48:22 2018

@author: 73902
"""

import pandas as pd
df=pd.read_excel('day.xlsx')
tran=df['tran']
sen=df['sentence']
ee=df['ee']
listge=[]

import xlrd

import xlwt


#doc
from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
excel_data = xlrd.open_workbook(r'd:\lucas\day.xlsx')
sheet = excel_data.sheets()[0]

#workbook=xlwt.Workbook()
#sheet=workbook.add_sheet('Sheet 1')
document = Document()
table = document.add_table(rows=50,cols=3)

for i in range(len(tran)):
    listge.append(str(tran[i]))
    listge.append('\n')
    listge.append(str(ee[i]))
    listge.append('\n')
    listge.append(str(sen[i]))
    print(listge)
#    sheet.write(i,0,listge)
    te=table.rows[i].cells
    te[2].text=listge
    listge=[]
    

document.save('demo.docx')    
#workbook.save("Python.xls")






tb=doc.tables
for i in range(2,41):
    
rowsw=tb[0].rows
cols=rowsw[0].cells
